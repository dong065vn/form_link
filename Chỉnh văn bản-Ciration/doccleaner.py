#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
DocCleaner GUI (Tkinter) — Enhanced Version
- Comprehensive document cleaning and processing tool
- Modular architecture with error handling
- Dark/Light theme support
- Multiple file format support (DOCX, TXT, MD, HTML)
- Citation management, translation, and analysis features

Requirements:
    pip install python-docx lxml beautifulsoup4
"""

import os
import re
import csv
import json
import zipfile
import logging
from pathlib import Path
from collections import Counter
from typing import Optional, List, Tuple, Dict
from tkinter import Tk, ttk, filedialog, messagebox, Text, END, StringVar, BooleanVar, Scrollbar

try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Inches
except ImportError:
    Document = None

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Stop words for filtering
STOP_VI = {
    'và', 'là', 'của', 'các', 'những', 'một', 'được', 'trong', 'cho', 'với',
    'đến', 'khi', 'này', 'đó', 'theo', 'từ', 'như', 'đã', 'sẽ', 'không',
    'cũng', 'rất', 'vì', 'do', 'ở', 'trên', 'dưới', 'giữa', 'nếu', 'thì',
    'vào', 'ra', 'lại', 'có', 'bởi', 'về', 'sau', 'trước', 'nữa', 'mà'
}

STOP_EN = {
    'the', 'a', 'an', 'and', 'or', 'but', 'for', 'to', 'of', 'in', 'on',
    'at', 'by', 'with', 'from', 'as', 'is', 'are', 'was', 'were', 'be',
    'been', 'it', 'this', 'that', 'these', 'those', 'not', 'can', 'will',
    'would', 'should', 'could', 'have', 'has', 'had', 'do', 'does', 'did'
}

# Translation dictionaries
BASE_EN = {
    'introduction': 'giới thiệu', 'conclusion': 'kết luận', 'summary': 'tóm tắt',
    'chapter': 'chương', 'references': 'tài liệu tham khảo', 'figure': 'hình',
    'table': 'bảng', 'method': 'phương pháp', 'result': 'kết quả',
    'discussion': 'thảo luận', 'transportation': 'giao thông',
    'wireless sensor network': 'mạng cảm biến không dây',
    'smart city': 'đô thị thông minh', 'abstract': 'tóm tắt',
    'keywords': 'từ khóa', 'acknowledgment': 'lời cảm ơn'
}

BASE_JA = {
    'はじめに': 'mở đầu', '結論': 'kết luận', '参考文献': 'tài liệu tham khảo',
    '図': 'hình', '表': 'bảng', '方法': 'phương pháp', '結果': 'kết quả',
    '考察': 'thảo luận', '交通': 'giao thông',
    '無線センサネットワーク': 'mạng cảm biến không dây',
    'スマートシティ': 'đô thị thông minh', '要約': 'tóm tắt'
}


class DocCleanerGUI:
    """Main GUI application for document cleaning and processing."""

    def __init__(self, root: Tk):
        """Initialize the DocCleaner GUI application."""
        self.root = root
        self.root.title('DocCleaner — Enhanced GUI')
        self.root.geometry('1280x800')
        
        # Set minimum window size
        self.root.minsize(900, 600)

        # State variables
        self.current_path: str = ''
        self.current_kind: str = ''
        self.refs: List[str] = []
        self.custom_dict: Dict[str, str] = {}
        self.undo_stack: List[str] = []
        self.max_undo: int = 20

        # UI state variables
        self.var_status = StringVar(value='Sẵn sàng.')
        self.var_quotes = StringVar(value='vni')
        self.var_justify = BooleanVar(value=False)
        self.var_dark = BooleanVar(value=True)
        self.var_word_count = StringVar(value='Từ: 0')

        self._build_styles()
        self._build_layout()
        self._apply_theme()
        
        logger.info("DocCleaner GUI initialized successfully")

    # ===== Theme Management =====
    def _build_styles(self) -> None:
        """Configure application styles and themes."""
        self.style = ttk.Style()
        
        # Use a modern theme if available
        available_themes = self.style.theme_names()
        for preferred_theme in ['clam', 'alt', 'default']:
            if preferred_theme in available_themes:
                try:
                    self.style.theme_use(preferred_theme)
                    break
                except Exception as e:
                    logger.warning(f"Could not set theme {preferred_theme}: {e}")

        # Light theme colors
        self.colors_light = {
            'bg': '#f6f7fb',
            'fg': '#111',
            'panel': '#ffffff',
            'muted': '#666',
            'accent': '#2563eb',
            'entrybg': '#ffffff',
            'textbg': '#ffffff',
            'textfg': '#111',
            'border': '#d1d5db'
        }
        
        # Dark theme colors
        self.colors_dark = {
            'bg': '#0b1020',
            'fg': '#e6e9f8',
            'panel': '#121a33',
            'muted': '#9aa4bf',
            'accent': '#6aa6ff',
            'entrybg': '#0e1530',
            'textbg': '#0e1530',
            'textfg': '#e6e9f8',
            'border': '#1e2944'
        }

    def _apply_theme(self) -> None:
        """Apply the selected theme (dark or light) to all UI elements."""
        c = self.colors_dark if self.var_dark.get() else self.colors_light
        
        # Root window
        self.root.configure(bg=c['bg'])
        
        # Frame styles
        for cls in ('TFrame', 'TLabelframe', 'TLabelframe.Label'):
            self.style.configure(cls, background=c['bg'], foreground=c['fg'])
        
        # Notebook styles
        self.style.configure('TNotebook', background=c['bg'], borderwidth=0)
        self.style.configure('TNotebook.Tab', 
                           background=c['panel'], 
                           foreground=c['fg'],
                           padding=[10, 5])
        self.style.map('TNotebook.Tab',
                      background=[('selected', c['accent'])],
                      foreground=[('selected', '#ffffff' if self.var_dark.get() else '#ffffff')])
        
        # Button styles
        self.style.configure('TButton', 
                           background=c['panel'], 
                           foreground=c['fg'],
                           borderwidth=1,
                           relief='flat')
        self.style.map('TButton',
                      background=[('active', c['accent'])],
                      foreground=[('active', '#ffffff')])
        
        # Label and Checkbutton
        self.style.configure('TLabel', background=c['bg'], foreground=c['fg'])
        self.style.configure('TCheckbutton', background=c['bg'], foreground=c['fg'])
        
        # Combobox
        self.style.configure('TCombobox',
                           fieldbackground=c['entrybg'],
                           background=c['panel'],
                           foreground=c['fg'])
        
        # Text widget (manual configuration)
        try:
            self.txt.configure(
                bg=c['textbg'],
                fg=c['textfg'],
                insertbackground=c['textfg'],
                selectbackground=c['accent'],
                selectforeground='#ffffff'
            )
        except AttributeError:
            pass  # Text widget not yet created

    # ===== Helper Methods =====
    def set_status(self, msg: str) -> None:
        """Update status bar message."""
        self.var_status.set(msg)
        self.root.update_idletasks()
        logger.info(f"Status: {msg}")

    def update_word_count(self) -> None:
        """Update word count in status bar."""
        try:
            text = self.txt.get('1.0', END).strip()
            words = len(re.findall(r'\S+', text))
            chars = len(text)
            self.var_word_count.set(f'Từ: {words:,} | Ký tự: {chars:,}')
        except Exception as e:
            logger.error(f"Error updating word count: {e}")

    def push_undo(self) -> None:
        """Push current text state to undo stack."""
        try:
            current_text = self.txt.get('1.0', END)
            self.undo_stack.append(current_text)
            
            # Limit undo stack size
            if len(self.undo_stack) > self.max_undo:
                self.undo_stack.pop(0)
        except Exception as e:
            logger.error(f"Error pushing to undo stack: {e}")

    def undo_action(self) -> None:
        """Undo last text modification."""
        if not self.undo_stack:
            messagebox.showinfo('Hoàn tác', 'Không có thao tác để hoàn tác.')
            return
        
        try:
            previous_text = self.undo_stack.pop()
            self.txt.delete('1.0', END)
            self.txt.insert('1.0', previous_text)
            self.set_status('Đã hoàn tác.')
            self.update_word_count()
        except Exception as e:
            logger.error(f"Error during undo: {e}")
            messagebox.showerror('Lỗi', f'Không thể hoàn tác: {e}')

    def read_text_from_file(self, path: str) -> Tuple[str, str]:
        """
        Read text content from various file formats.
        
        Args:
            path: File path to read from
            
        Returns:
            Tuple of (text_content, file_type)
        """
        try:
            ext = os.path.splitext(path)[1].lower()
            
            if ext == '.docx':
                if Document is None:
                    messagebox.showerror('Lỗi', 
                        'Thiếu thư viện python-docx.\nChạy: pip install python-docx')
                    return '', ''
                doc = Document(path)
                text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
                return text, 'docx'
            
            elif ext in {'.txt', '.md'}:
                with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read(), 'text'
            
            elif ext in {'.html', '.htm'}:
                with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                    html = f.read()
                
                if BeautifulSoup:
                    soup = BeautifulSoup(html, 'html.parser')
                    # Remove script and style elements
                    for script in soup(['script', 'style']):
                        script.decompose()
                    text = soup.get_text('\n')
                else:
                    text = re.sub(r'<[^>]+>', '', html)
                
                # Clean up whitespace
                text = re.sub(r'\n\s*\n', '\n\n', text)
                return text.strip(), 'html'
            
            else:
                messagebox.showwarning('Chưa hỗ trợ', 
                    f'Định dạng {ext} chưa được hỗ trợ.')
                return '', ''
                
        except Exception as e:
            logger.error(f"Error reading file {path}: {e}")
            messagebox.showerror('Lỗi', f'Không thể đọc tệp:\n{e}')
            return '', ''

    def _escape_html(self, s: str) -> str:
        """Escape HTML special characters."""
        return (s.replace('&', '&amp;')
                 .replace('<', '&lt;')
                 .replace('>', '&gt;')
                 .replace('"', '&quot;')
                 .replace("'", '&#39;'))

    # ===== Text Processing Operations =====
    def normalize_quotes(self, text: str, mode: str = 'vni') -> str:
        """
        Normalize quotation marks based on style.
        
        Args:
            text: Input text
            mode: 'vni' for Vietnamese style or 'ieee' for standard quotes
        """
        if mode == 'vni':
            # Vietnamese style: "" and ''
            text = re.sub(r'"([^"]*)"', '\u201c\\1\u201d', text)
            text = re.sub(r"'([^']*)'", '\u2018\\1\u2019', text)
        else:
            # Standard style: "" and ''
            text = re.sub(r'[\u201c\u201d]', '"', text)
            text = re.sub(r'[\u2018\u2019]', "'", text)
        return text

    def citations_normalize(self, text: str) -> str:
        """Normalize citation formats."""
        # Normalize numbered citations [1], [2, 3]
        text = re.sub(r'\[\s*(\d+)\s*\]', r'[\1]', text)
        text = re.sub(r'\[\s*(\d+)\s*,\s*(\d+)\s*\]', r'[\1, \2]', text)
        
        # Normalize author-year citations (Smith, 2020)
        text = re.sub(r'\(\s*([A-Z][A-Za-z\-]+(?:\s+et\s+al\.)?)\s*,\s*(\d{4}[a-z]?)\s*\)',
                     r'(\1, \2)', text)
        
        return text

    def citations_remove(self, text: str) -> str:
        """Remove all citations from text."""
        # Remove numbered citations [1], [1-3], [1, 2, 3]
        text = re.sub(r'\[(?:\s*\d+\s*(?:[–\-,]\s*\d+\s*)*)+\]', '', text)
        
        # Remove author-year citations (Smith, 2020), (Smith & Jones, 2020)
        text = re.sub(r'\([A-Z][A-Za-z\-]+(?:\s+et\s+al\.)?'
                     r'(?:\s*&\s*[A-Z][A-Za-z\-]+)?,\s*\d{4}[a-z]?\)', '', text)
        
        # Clean up extra spaces
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\s+([.,;:])', r'\1', text)
        
        return text

    def links_strip_and_collect(self, text: str) -> str:
        """Remove links from text and collect them for reference."""
        # Find all URLs
        urls = re.findall(r'https?://[^\s<>"\')]+', text)
        
        # Remove HTML anchor tags but keep link text
        text = re.sub(r'<a[^>]*>([\s\S]*?)</a>', r'\1', text, flags=re.IGNORECASE)
        
        # Remove plain URLs
        text = re.sub(r'https?://[^\s<>"\')]+', '', text)
        
        # Clean up extra spaces
        text = re.sub(r'\s+', ' ', text)
        
        # Store unique URLs
        self.refs = list(dict.fromkeys(urls))
        
        return text

    def scan_issues(self, text: str) -> List[str]:
        """Scan text for common formatting issues."""
        issues = []
        
        # Double periods
        double_dots = re.findall(r'\.\.', text)
        if double_dots:
            issues.append(f'Dấu câu: {len(double_dots)} vị trí có ".."')
        
        # Space before punctuation
        bad_space = re.findall(r'\s+[.,;:]', text)
        if bad_space:
            issues.append(f'Khoảng trắng: {len(bad_space)} vị trí có dấu cách trước dấu câu')
        
        # Missing capitalization at sentence start
        bad_cap = re.findall(r'(?:^|[.!?]\s+)[a-zà-ỹ]', text, re.MULTILINE)
        if bad_cap:
            issues.append(f'Viết hoa: {len(bad_cap)} câu có thể thiếu viết hoa đầu câu')
        
        # Multiple spaces
        multi_space = re.findall(r'  +', text)
        if multi_space:
            issues.append(f'Khoảng trắng: {len(multi_space)} vị trí có nhiều dấu cách liên tiếp')
        
        # Missing space after punctuation
        no_space = re.findall(r'[.,;:][A-ZÀ-Ỹa-zà-ỹ]', text)
        if no_space:
            issues.append(f'Khoảng trắng: {len(no_space)} vị trí thiếu dấu cách sau dấu câu')
        
        return issues

    def keyword_stats(self, text: str, topk: int = 50) -> List[Tuple[str, int]]:
        """Generate keyword statistics from text."""
        t = text.lower()
        
        # Extract words (alphanumeric and Vietnamese characters)
        words = re.findall(r'[a-zà-ỹ0-9]+', t)
        
        # Detect language (Vietnamese or English)
        vi_count = sum(1 for w in words if re.search(r'[à-ỹ]', w))
        stop = STOP_VI if vi_count > len(words) / 10 else STOP_EN
        
        # Filter stop words and short words
        filtered = [w for w in words if w not in stop and len(w) >= 3]
        
        return Counter(filtered).most_common(topk)

    def toc_detect(self, text: str) -> List[Tuple[int, str]]:
        """Detect table of contents entries."""
        results = []
        
        for i, line in enumerate(text.splitlines(), 1):
            s = line.strip()
            
            # Match chapter headers: "CHƯƠNG 1", "1.", "1.1", etc.
            if re.match(r'^(CHƯƠNG\s+[IVX\d]+|\d+\.(?:\d+\.)*\s+)', s, flags=re.IGNORECASE):
                results.append((i, s[:150]))
            
            # Match common section headers
            elif re.match(r'^(Abstract|Introduction|Conclusion|References|Acknowledgment)', 
                         s, flags=re.IGNORECASE):
                results.append((i, s[:150]))
        
        return results

    def translate_simple(self, text: str, src: str, 
                        glossary_path: Optional[str]) -> str:
        """Simple dictionary-based translation."""
        # Select base dictionary
        d = BASE_JA.copy() if src == 'ja' else BASE_EN.copy()
        
        # Load custom glossary if provided
        if glossary_path and os.path.exists(glossary_path):
            try:
                with open(glossary_path, 'r', encoding='utf-8') as f:
                    custom = json.load(f)
                    d.update(custom)
                logger.info(f"Loaded {len(custom)} custom translations")
            except Exception as e:
                logger.error(f"Error loading glossary: {e}")
        
        # Apply translations (longest matches first)
        for key in sorted(d.keys(), key=len, reverse=True):
            pattern = re.escape(key)
            text = re.sub(pattern, d[key], text, flags=re.IGNORECASE)
        
        return text

    def apply_formatting_docx(self, in_path: str, out_path: str, 
                             justify: bool) -> None:
        """Apply formatting to DOCX file."""
        if Document is None:
            messagebox.showerror('Lỗi', 
                'Thiếu thư viện python-docx.\nChạy: pip install python-docx')
            return
        
        try:
            doc = Document(in_path)
            
            for paragraph in doc.paragraphs:
                # Set alignment
                if justify:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # Format runs
                for run in paragraph.runs:
                    try:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                    except Exception as e:
                        logger.warning(f"Could not format run: {e}")
                
                # Set line spacing
                try:
                    paragraph.paragraph_format.line_spacing = 1.5
                    paragraph.paragraph_format.space_after = Pt(6)
                except Exception as e:
                    logger.warning(f"Could not set spacing: {e}")
            
            doc.save(out_path)
            logger.info(f"Formatted DOCX saved to {out_path}")
            
        except Exception as e:
            logger.error(f"Error formatting DOCX: {e}")
            messagebox.showerror('Lỗi', f'Không thể định dạng DOCX:\n{e}')

    def remove_docx_watermark(self, in_path: str, out_path: str) -> None:
        """Remove watermarks from DOCX file."""
        try:
            with zipfile.ZipFile(in_path, 'r') as zin:
                with zipfile.ZipFile(out_path, 'w', zipfile.ZIP_DEFLATED) as zout:
                    for item in zin.infolist():
                        data = zin.read(item.filename)
                        
                        # Process XML files that may contain watermarks
                        if re.search(r'word/(header\d+\.xml|footer\d+\.xml|document\.xml)$', 
                                   item.filename):
                            try:
                                xml = data.decode('utf-8', errors='ignore')
                                
                                # Remove picture elements
                                xml = re.sub(r'<w:pict[\s\S]*?</w:pict>', '', xml)
                                xml = re.sub(r'<v:shape[\s\S]*?</v:shape>', '', xml)
                                
                                # Remove drawings that look like watermarks
                                def filter_drawing(match):
                                    block = match.group(0)
                                    if re.search(r'watermark|mso-text-shadow|opacity',
                                               block, flags=re.IGNORECASE):
                                        return ''
                                    return block
                                
                                xml = re.sub(r'<w:drawing[\s\S]*?</w:drawing>', 
                                           filter_drawing, xml)
                                
                                data = xml.encode('utf-8')
                            except Exception as e:
                                logger.warning(f"Could not process {item.filename}: {e}")
                        
                        zout.writestr(item, data)
            
            logger.info(f"Watermark removed, saved to {out_path}")
            
        except Exception as e:
            logger.error(f"Error removing watermark: {e}")
            messagebox.showerror('Lỗi', f'Không thể gỡ watermark:\n{e}')

    # ===== UI Construction =====
    def _build_layout(self) -> None:
        """Build the main UI layout."""
        # Top toolbar
        top = ttk.Frame(self.root, padding=8)
        top.pack(fill='x')

        ttk.Button(top, text='📂 Mở tệp', command=self.ui_open).pack(side='left', padx=2)
        ttk.Button(top, text='💾 Lưu HTML', command=self.ui_export_html).pack(side='left', padx=2)
        ttk.Button(top, text='💾 Lưu DOCX', command=self.ui_export_docx).pack(side='left', padx=2)
        ttk.Button(top, text='↩️ Hoàn tác', command=self.undo_action).pack(side='left', padx=2)
        
        ttk.Label(top, textvariable=self.var_word_count, 
                 foreground='gray').pack(side='right', padx=10)
        ttk.Checkbutton(top, text='🌙 Dark Mode', 
                       variable=self.var_dark, 
                       command=self._apply_theme).pack(side='right', padx=6)
        ttk.Label(top, textvariable=self.var_status).pack(side='right', padx=10)

        # Main content area with notebook
        nb = ttk.Notebook(self.root)
        nb.pack(fill='both', expand=True, padx=8, pady=8)

        # Create tabs
        self.tab_file = ttk.Frame(nb, padding=10)
        self.tab_clean = ttk.Frame(nb, padding=10)
        self.tab_cite = ttk.Frame(nb, padding=10)
        self.tab_trans = ttk.Frame(nb, padding=10)
        self.tab_analysis = ttk.Frame(nb, padding=10)
        self.tab_docx = ttk.Frame(nb, padding=10)
        self.tab_export = ttk.Frame(nb, padding=10)

        nb.add(self.tab_file, text='📄 Tệp')
        nb.add(self.tab_clean, text='🧹 Làm sạch')
        nb.add(self.tab_cite, text='📚 Trích dẫn & Link')
        nb.add(self.tab_trans, text='🌐 Dịch')
        nb.add(self.tab_analysis, text='📊 Phân tích')
        nb.add(self.tab_docx, text='📝 DOCX & Watermark')
        nb.add(self.tab_export, text='💾 Xuất')

        self._build_file_tab()
        self._build_clean_tab()
        self._build_cite_tab()
        self._build_trans_tab()
        self._build_analysis_tab()
        self._build_docx_tab()
        self._build_export_tab()

    def _build_file_tab(self) -> None:
        """Build file preview tab."""
        # Instructions
        info = ttk.Label(self.tab_file, 
                        text='Mở tệp để xem và chỉnh sửa nội dung. '
                             'Hỗ trợ: DOCX, TXT, MD, HTML')
        info.pack(anchor='w', pady=(0, 10))
        
        # Text area with scrollbar
        text_frame = ttk.Frame(self.tab_file)
        text_frame.pack(fill='both', expand=True)
        
        scrollbar = Scrollbar(text_frame)
        scrollbar.pack(side='right', fill='y')
        
        self.txt = Text(text_frame, wrap='word', 
                       yscrollcommand=scrollbar.set,
                       font=('Consolas', 10))
        self.txt.pack(side='left', fill='both', expand=True)
        scrollbar.config(command=self.txt.yview)
        
        # Bind text change event for word count
        self.txt.bind('<<Modified>>', lambda e: self.update_word_count())

    def _build_clean_tab(self) -> None:
        """Build cleaning operations tab."""
        ttk.Label(self.tab_clean, 
                 text='Các công cụ làm sạch và chuẩn hóa văn bản',
                 font=('TkDefaultFont', 10, 'bold')).pack(anchor='w', pady=(0, 10))
        
        ttk.Button(self.tab_clean, text='🧹 Lọc ký tự ẩn/thừa', 
                  command=self.ui_clean_hidden).pack(anchor='w', pady=4, fill='x')
        
        ttk.Button(self.tab_clean, text='📄 Xóa dòng tên file', 
                  command=self.ui_remove_filename_lines).pack(anchor='w', pady=4, fill='x')
        
        ttk.Button(self.tab_clean, text='✨ Chuẩn hóa khoảng trắng', 
                  command=self.ui_normalize_whitespace).pack(anchor='w', pady=4, fill='x')
        
        ttk.Separator(self.tab_clean, orient='horizontal').pack(fill='x', pady=10)
        
        # Quote normalization
        quote_frame = ttk.LabelFrame(self.tab_clean, text='Chuẩn hóa ngoặc', padding=10)
        quote_frame.pack(fill='x', pady=5)
        
        row = ttk.Frame(quote_frame)
        row.pack(fill='x')
        ttk.Label(row, text='Phong cách:').pack(side='left', padx=(0, 10))
        ttk.Combobox(row, textvariable=self.var_quotes, 
                    values=['vni', 'ieee'], width=10, 
                    state='readonly').pack(side='left')
        ttk.Button(quote_frame, text='Áp dụng', 
                  command=self.ui_quotes).pack(anchor='w', pady=(10, 0))

    def _build_cite_tab(self) -> None:
        """Build citations and links tab."""
        ttk.Label(self.tab_cite, 
                 text='Quản lý trích dẫn và liên kết',
                 font=('TkDefaultFont', 10, 'bold')).pack(anchor='w', pady=(0, 10))
        
        ttk.Button(self.tab_cite, text='📝 Chuẩn hóa citation', 
                  command=self.ui_cite_norm).pack(anchor='w', pady=4, fill='x')
        
        ttk.Button(self.tab_cite, text='🗑️ Xóa citation', 
                  command=self.ui_cite_remove).pack(anchor='w', pady=4, fill='x')
        
        ttk.Separator(self.tab_cite, orient='horizontal').pack(fill='x', pady=10)
        
        ttk.Button(self.tab_cite, text='🔗 Xóa link & tách references', 
                  command=self.ui_links_strip).pack(anchor='w', pady=4, fill='x')

    def _build_trans_tab(self) -> None:
        """Build translation tab."""
        ttk.Label(self.tab_trans, 
                 text='Dịch văn bản đơn giản (dựa trên từ điển)',
                 font=('TkDefaultFont', 10, 'bold')).pack(anchor='w', pady=(0, 10))
        
        ttk.Button(self.tab_trans, text='🇬🇧 Dịch EN → VI', 
                  command=lambda: self.ui_translate('en')).pack(anchor='w', pady=4, fill='x')
        
        ttk.Button(self.tab_trans, text='🇯🇵 Dịch JA → VI', 
                  command=lambda: self.ui_translate('ja')).pack(anchor='w', pady=4, fill='x')
        
        ttk.Separator(self.tab_trans, orient='horizontal').pack(fill='x', pady=10)
        
        ttk.Button(self.tab_trans, text='📖 Nạp glossary tùy chỉnh', 
                  command=self.ui_load_glossary).pack(anchor='w', pady=4, fill='x')

    def _build_analysis_tab(self) -> None:
        """Build analysis tab."""
        ttk.Label(self.tab_analysis, 
                 text='Phân tích và thống kê văn bản',
                 font=('TkDefaultFont', 10, 'bold')).pack(anchor='w', pady=(0, 10))
        
        ttk.Button(self.tab_analysis, text='🔍 Quét lỗi định dạng', 
                  command=self.ui_scan).pack(anchor='w', pady=4, fill='x')
        
        ttk.Button(self.tab_analysis, text='📊 Thống kê từ khóa → CSV', 
                  command=self.ui_keywords).pack(anchor='w', pady=4, fill='x')
        
        ttk.Button(self.tab_analysis, text='📑 Tạo mục lục tự động', 
                  command=self.ui_toc).pack(anchor='w', pady=4, fill='x')

    def _build_docx_tab(self) -> None:
        """Build DOCX operations tab."""
        ttk.Label(self.tab_docx, 
                 text='Định dạng và xử lý file DOCX',
                 font=('TkDefaultFont', 10, 'bold')).pack(anchor='w', pady=(0, 10))
        
        # Formatting options
        format_frame = ttk.LabelFrame(self.tab_docx, text='Định dạng', padding=10)
        format_frame.pack(fill='x', pady=5)
        
        ttk.Checkbutton(format_frame, text='Căn lề hai bên (Justify)', 
                       variable=self.var_justify).pack(anchor='w')
        
        ttk.Button(format_frame, text='Áp dụng định dạng (TNR 12, Line 1.5)', 
                  command=self.ui_format_docx).pack(anchor='w', pady=(10, 0), fill='x')
        
        ttk.Separator(self.tab_docx, orient='horizontal').pack(fill='x', pady=10)
        
        # Watermark removal
        ttk.Button(self.tab_docx, text='💧 Gỡ watermark DOCX', 
                  command=self.ui_nowm).pack(anchor='w', pady=4, fill='x')

    def _build_export_tab(self) -> None:
        """Build export tab."""
        ttk.Label(self.tab_export, 
                 text='Xuất văn bản sang các định dạng khác',
                 font=('TkDefaultFont', 10, 'bold')).pack(anchor='w', pady=(0, 10))
        
        ttk.Button(self.tab_export, text='🌐 Xuất HTML (học thuật)', 
                  command=self.ui_export_html).pack(anchor='w', pady=4, fill='x')
        
        ttk.Button(self.tab_export, text='📝 Xuất DOCX (từ nội dung hiện tại)', 
                  command=self.ui_export_docx).pack(anchor='w', pady=4, fill='x')

    # ===== UI Event Handlers =====
    def ui_open(self) -> None:
        """Handle file open action."""
        path = filedialog.askopenfilename(
            title='Chọn tệp để mở',
            filetypes=[
                ('Tất cả hỗ trợ', '*.docx;*.txt;*.md;*.html;*.htm'),
                ('DOCX', '*.docx'),
                ('Text', '*.txt;*.md'),
                ('HTML', '*.html;*.htm'),
                ('Tất cả', '*.*')
            ]
        )
        
        if not path:
            return
        
        self.set_status('Đang mở tệp...')
        text, kind = self.read_text_from_file(path)
        
        if not kind:
            return
        
        self.current_path = path
        self.current_kind = kind
        self.undo_stack.clear()
        
        self.txt.delete('1.0', END)
        self.txt.insert('1.0', text)
        
        self.update_word_count()
        self.set_status(f'✅ Đã mở: {os.path.basename(path)}')

    def _replace_text(self, text: str) -> None:
        """Replace text content with undo support."""
        self.push_undo()
        self.txt.delete('1.0', END)
        self.txt.insert('1.0', text)
        self.update_word_count()

    def ui_clean_hidden(self) -> None:
        """Clean hidden and invalid characters."""
        text = self.txt.get('1.0', END)
        
        # Remove control characters
        text = re.sub(r'[\x00-\x1F\x7F-\x9F]', '', text)
        
        # Remove zero-width characters
        text = re.sub(r'[\u200B-\u200D\u2060\uFEFF]', '', text)
        
        # Replace non-breaking space with regular space
        text = text.replace('\u00A0', ' ')
        
        # Remove trailing whitespace
        text = re.sub(r'[ \t]+$', '', text, flags=re.MULTILINE)
        
        # Normalize multiple blank lines
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        self._replace_text(text.strip())
        self.set_status('✅ Đã lọc ký tự ẩn/thừa')

    def ui_normalize_whitespace(self) -> None:
        """Normalize all whitespace in text."""
        text = self.txt.get('1.0', END)
        
        # Remove space before punctuation
        text = re.sub(r'\s+([.,;:!?])', r'\1', text)
        
        # Add space after punctuation if missing
        text = re.sub(r'([.,;:!?])([A-ZÀ-Ỹa-zà-ỹ])', r'\1 \2', text)
        
        # Normalize multiple spaces to single space
        text = re.sub(r' {2,}', ' ', text)
        
        # Normalize line breaks
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        self._replace_text(text.strip())
        self.set_status('✅ Đã chuẩn hóa khoảng trắng')

    def ui_remove_filename_lines(self) -> None:
        """Remove lines that are just filenames."""
        text = self.txt.get('1.0', END)
        pattern = re.compile(
            r'(?i)^[\w\-\s]+\.(docx?|pdf|pptx?|xlsx?|csv|txt|md|html?|'
            r'jpg|jpeg|png|gif|bmp|svg|zip|rar|7z|tar|gz)$'
        )
        
        lines = []
        removed = 0
        
        for line in text.splitlines():
            stripped = line.strip()
            if pattern.match(stripped):
                removed += 1
                continue
            lines.append(line)
        
        self._replace_text('\n'.join(lines))
        self.set_status(f'✅ Đã xóa {removed} dòng tên file')

    def ui_quotes(self) -> None:
        """Normalize quotation marks."""
        text = self.txt.get('1.0', END)
        text = self.normalize_quotes(text, self.var_quotes.get())
        self._replace_text(text)
        self.set_status(f'✅ Đã chuẩn hóa ngoặc ({self.var_quotes.get()})')

    def ui_cite_norm(self) -> None:
        """Normalize citations."""
        text = self.txt.get('1.0', END)
        text = self.citations_normalize(text)
        self._replace_text(text)
        self.set_status('✅ Đã chuẩn hóa citation')

    def ui_cite_remove(self) -> None:
        """Remove all citations."""
        if not messagebox.askyesno('Xác nhận', 
                                   'Xóa tất cả citation?\nHành động này không thể hoàn tác.'):
            return
        
        text = self.txt.get('1.0', END)
        text = self.citations_remove(text)
        self._replace_text(text)
        self.set_status('✅ Đã xóa citation')

    def ui_links_strip(self) -> None:
        """Strip links and collect references."""
        text = self.txt.get('1.0', END)
        text = self.links_strip_and_collect(text)
        self._replace_text(text)
        
        if self.refs:
            save_path = filedialog.asksaveasfilename(
                title='Lưu danh sách references',
                defaultextension='.txt',
                filetypes=[('Text', '*.txt'), ('All', '*.*')]
            )
            
            if save_path:
                try:
                    with open(save_path, 'w', encoding='utf-8') as f:
                        f.write('# References\n\n')
                        for i, url in enumerate(self.refs, 1):
                            f.write(f'{i}. {url}\n')
                    self.set_status(f'✅ Đã tách {len(self.refs)} URL → {os.path.basename(save_path)}')
                except Exception as e:
                    messagebox.showerror('Lỗi', f'Không thể lưu references:\n{e}')
        else:
            self.set_status('✅ Đã xóa link (không tìm thấy URL)')

    def ui_scan(self) -> None:
        """Scan for common issues."""
        text = self.txt.get('1.0', END)
        issues = self.scan_issues(text)
        
        if not issues:
            messagebox.showinfo('Quét lỗi', 
                              '✅ Không phát hiện lỗi định dạng thường gặp!')
        else:
            msg = '⚠️ Phát hiện các vấn đề:\n\n' + '\n'.join(f'• {issue}' for issue in issues)
            messagebox.showwarning('Quét lỗi', msg)

    def ui_keywords(self) -> None:
        """Extract and export keywords."""
        text = self.txt.get('1.0', END)
        
        if not text.strip():
            messagebox.showwarning('Cảnh báo', 'Không có nội dung để phân tích.')
            return
        
        data = self.keyword_stats(text, topk=50)
        
        save_path = filedialog.asksaveasfilename(
            title='Lưu thống kê từ khóa',
            defaultextension='.csv',
            filetypes=[('CSV', '*.csv'), ('All', '*.*')]
        )
        
        if not save_path:
            return
        
        try:
            with open(save_path, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                writer.writerow(['Từ khóa', 'Số lần'])
                for keyword, count in data:
                    writer.writerow([keyword, count])
            
            self.set_status(f'✅ Đã lưu {len(data)} từ khóa → {os.path.basename(save_path)}')
            
        except Exception as e:
            logger.error(f"Error saving keywords: {e}")
            messagebox.showerror('Lỗi', f'Không thể lưu CSV:\n{e}')

    def ui_toc(self) -> None:
        """Detect and display table of contents."""
        text = self.txt.get('1.0', END)
        toc = self.toc_detect(text)
        
        if not toc:
            messagebox.showinfo('Mục lục', 
                              'Không nhận diện được tiêu đề chương/phần.')
            return
        
        msg = '📑 Mục lục tạm:\n\n' + '\n'.join(
            f'Dòng {line}: {title}' for line, title in toc
        )
        
        messagebox.showinfo('Mục lục', msg)

    def ui_translate(self, src: str) -> None:
        """Translate text using dictionary."""
        text = self.txt.get('1.0', END)
        
        if not text.strip():
            messagebox.showwarning('Cảnh báo', 'Không có nội dung để dịch.')
            return
        
        # Ask for glossary
        glossary = filedialog.askopenfilename(
            title='Chọn glossary JSON (tùy chọn, bỏ qua nếu không có)',
            filetypes=[('JSON', '*.json'), ('All', '*.*')]
        )
        
        self.set_status(f'Đang dịch {src.upper()} → VI...')
        translated = self.translate_simple(text, src, glossary if glossary else None)
        
        self._replace_text(translated)
        self.set_status(f'✅ Đã dịch {src.upper()} → VI')

    def ui_load_glossary(self) -> None:
        """Load custom glossary."""
        path = filedialog.askopenfilename(
            title='Chọn glossary JSON',
            filetypes=[('JSON', '*.json'), ('All', '*.*')]
        )
        
        if not path:
            return
        
        try:
            with open(path, 'r', encoding='utf-8') as f:
                self.custom_dict = json.load(f)
            
            messagebox.showinfo('Glossary', 
                              f'✅ Đã nạp {len(self.custom_dict)} mục từ glossary')
            logger.info(f"Loaded glossary with {len(self.custom_dict)} entries")
            
        except Exception as e:
            logger.error(f"Error loading glossary: {e}")
            messagebox.showerror('Lỗi', f'Không thể nạp glossary:\n{e}')

    def ui_format_docx(self) -> None:
        """Format DOCX file."""
        if not self.current_path or \
           os.path.splitext(self.current_path)[1].lower() != '.docx':
            messagebox.showwarning('Cảnh báo', 
                                 'Hãy mở một tệp DOCX trước khi định dạng.')
            return
        
        save_path = filedialog.asksaveasfilename(
            title='Lưu DOCX đã định dạng',
            defaultextension='.docx',
            filetypes=[('DOCX', '*.docx')]
        )
        
        if not save_path:
            return
        
        self.set_status('Đang định dạng DOCX...')
        self.apply_formatting_docx(self.current_path, save_path, 
                                  self.var_justify.get())
        self.set_status(f'✅ Đã định dạng → {os.path.basename(save_path)}')

    def ui_nowm(self) -> None:
        """Remove watermark from DOCX."""
        path = filedialog.askopenfilename(
            title='Chọn DOCX để gỡ watermark',
            filetypes=[('DOCX', '*.docx'), ('All', '*.*')]
        )
        
        if not path:
            return
        
        save_path = filedialog.asksaveasfilename(
            title='Lưu DOCX (không watermark)',
            defaultextension='.docx',
            filetypes=[('DOCX', '*.docx')]
        )
        
        if not save_path:
            return
        
        self.set_status('Đang gỡ watermark...')
        self.remove_docx_watermark(path, save_path)
        self.set_status(f'✅ Đã gỡ watermark → {os.path.basename(save_path)}')

    def ui_export_html(self) -> None:
        """Export text as HTML."""
        text = self.txt.get('1.0', END).strip()
        
        if not text:
            messagebox.showwarning('Cảnh báo', 'Không có nội dung để xuất.')
            return
        
        save_path = filedialog.asksaveasfilename(
            title='Lưu HTML',
            defaultextension='.html',
            filetypes=[('HTML', '*.html'), ('All', '*.*')]
        )
        
        if not save_path:
            return
        
        try:
            # Generate HTML with academic styling
            paragraphs = re.split(r'\n\n+', text)
            body = '\n'.join(f'<p>{self._escape_html(p)}</p>' for p in paragraphs if p.strip())
            
            html = f"""<!DOCTYPE html>
<html lang="vi">
<head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>DocCleaner Export</title>
    <style>
        body {{
            font-family: 'Times New Roman', Times, serif;
            font-size: 12pt;
            line-height: 1.5;
            max-width: 800px;
            margin: 40px auto;
            padding: 20px;
            text-align: justify;
        }}
        p {{
            margin: 0 0 12pt 0;
            text-indent: 1.5em;
        }}
        @media print {{
            body {{ margin: 0; padding: 20mm; }}
        }}
    </style>
</head>
<body>
{body}
</body>
</html>"""
            
            with open(save_path, 'w', encoding='utf-8') as f:
                f.write(html)
            
            self.set_status(f'✅ Đã xuất HTML → {os.path.basename(save_path)}')
            
        except Exception as e:
            logger.error(f"Error exporting HTML: {e}")
            messagebox.showerror('Lỗi', f'Không thể xuất HTML:\n{e}')

    def ui_export_docx(self) -> None:
        """Export text as DOCX."""
        if Document is None:
            messagebox.showerror('Lỗi', 
                'Thiếu thư viện python-docx.\nChạy: pip install python-docx')
            return
        
        text = self.txt.get('1.0', END).strip()
        
        if not text:
            messagebox.showwarning('Cảnh báo', 'Không có nội dung để xuất.')
            return
        
        save_path = filedialog.asksaveasfilename(
            title='Lưu DOCX',
            defaultextension='.docx',
            filetypes=[('DOCX', '*.docx')]
        )
        
        if not save_path:
            return
        
        try:
            doc = Document()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Add paragraphs
            for para_text in re.split(r'\n\n+', text):
                if para_text.strip():
                    p = doc.add_paragraph(para_text.strip())
                    p.paragraph_format.line_spacing = 1.5
                    p.paragraph_format.space_after = Pt(6)
                    
                    if self.var_justify.get():
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
            
            doc.save(save_path)
            self.set_status(f'✅ Đã xuất DOCX → {os.path.basename(save_path)}')
            
        except Exception as e:
            logger.error(f"Error exporting DOCX: {e}")
            messagebox.showerror('Lỗi', f'Không thể xuất DOCX:\n{e}')


def main():
    """Main entry point."""
    try:
        root = Tk()
        app = DocCleanerGUI(root)
        root.mainloop()
    except Exception as e:
        logger.error(f"Fatal error: {e}", exc_info=True)
        messagebox.showerror('Lỗi nghiêm trọng', 
                           f'Ứng dụng gặp lỗi:\n{e}')


if __name__ == '__main__':
    main()
